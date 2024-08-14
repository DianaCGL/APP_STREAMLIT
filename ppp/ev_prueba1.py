import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np

# Definir las descripciones de las rúbricas específicas para cada pregunta
rubricas = {
    'Gestión de Acceso': {
        '¿Existen políticas y procedimientos documentados para la gestión de accesos?': {
            1: 'No se tienen políticas ni procedimientos documentados para la gestión de accesos. Esto implica que no hay controles formales para regular quién tiene acceso a qué información y sistemas...',
            2: 'Existen políticas y procedimientos para la gestión de accesos, pero no están completamente documentados o actualizados. Esto puede llevar a inconsistencias en su aplicación...',
            3: 'Políticas y procedimientos para la gestión de accesos están documentados y se revisan regularmente. La mayoría de los requisitos de la norma ISO 27001 están cubiertos...',
            4: 'Las políticas y procedimientos de gestión de accesos cumplen completamente con los requisitos establecidos por la norma ISO 27001...',
            5: 'La implementación de políticas y procedimientos de gestión de accesos no solo cumple con todos los requisitos de la norma ISO 27001, sino que también incluye controles adicionales...'
        },
        '¿Se implementan controles de autenticación fuertes para acceder a sistemas críticos?': {
            1: 'No se implementan controles de autenticación para acceder a sistemas críticos, lo que significa que cualquier persona podría potencialmente acceder a información y recursos sensibles sin ninguna verificación...',
            2: 'Se implementan controles de autenticación de manera limitada o inconsistente. Algunos sistemas críticos pueden tener autenticación básica, como contraseñas simples...',
            3: 'Controles de autenticación fuertes están implementados de manera regular, cubriendo la mayoría de los sistemas críticos...',
            4: 'Los controles de autenticación cumplen totalmente con los requisitos de autenticación de ISO 27001...',
            5: 'La implementación de controles de autenticación es avanzada y supera los requisitos estándar...'
        }
    },
    'Seguridad Física y Ambiental': {
        '¿Existen medidas de seguridad física para proteger los equipos críticos del departamento de sistemas?': {
            1: 'No hay medidas de seguridad física implementadas, dejando los equipos críticos vulnerables a accesos no autorizados y daños físicos...',
            2: 'Las medidas de seguridad física existen pero son parciales o insuficientes...',
            3: 'Las medidas de seguridad física están bien implementadas y protegen la mayoría de los equipos críticos...',
            4: 'Las medidas de seguridad física cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de medidas de seguridad física es avanzada y supera los requisitos estándar...'
        },
        '¿Se realizan controles ambientales para proteger la infraestructura tecnológica (temperatura, humedad, etc.)?': {
            1: 'No se realizan controles ambientales, lo que puede llevar a fallos en la infraestructura tecnológica debido a condiciones ambientales adversas...',
            2: 'Los controles ambientales existen pero se aplican de manera irregular o insuficiente...',
            3: 'Controles ambientales están bien implementados y protegen la mayoría de la infraestructura tecnológica...',
            4: 'Los controles ambientales cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de controles ambientales es avanzada y supera los requisitos estándar...'
        }
    },
    'Gestión de Comunicaciones y Operaciones': {
        '¿Se utilizan procedimientos seguros para la transmisión de datos sensibles dentro y fuera de la organización?': {
            1: 'No se utilizan procedimientos seguros para la transmisión de datos, lo que expone la información sensible a interceptaciones y accesos no autorizados...',
            2: 'Los procedimientos seguros para la transmisión de datos existen pero se utilizan de manera limitada o inconsistente...',
            3: 'Procedimientos seguros para la transmisión de datos están bien implementados y se utilizan regularmente...',
            4: 'Los procedimientos seguros para la transmisión de datos cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de procedimientos seguros para la transmisión de datos es avanzada y supera los requisitos estándar...'
        },
        '¿Se realizan pruebas periódicas de vulnerabilidades y evaluaciones de riesgos en la infraestructura de redes?': {
            1: 'No se realizan pruebas de vulnerabilidades ni evaluaciones de riesgos, lo que deja la infraestructura de redes expuesta a posibles amenazas y ataques...',
            2: 'Las pruebas de vulnerabilidades y evaluaciones de riesgos se realizan de manera limitada o irregular...',
            3: 'Las pruebas de vulnerabilidades y evaluaciones de riesgos se realizan regularmente y cubren la mayoría de la infraestructura de redes...',
            4: 'Las pruebas de vulnerabilidades y evaluaciones de riesgos cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de pruebas de vulnerabilidades y evaluaciones de riesgos es avanzada y supera los requisitos estándar...'
        }
    },
    'Control de Acceso a la Información': {
        '¿Se implementan controles para limitar el acceso a la información confidencial y crítica dentro del departamento de sistemas?': {
            1: 'No se implementan controles para limitar el acceso a la información confidencial y crítica, lo que significa que cualquier persona dentro del departamento puede acceder a estos datos sin restricciones...',
            2: 'Los controles de acceso a la información confidencial y crítica existen pero se implementan de manera limitada o inconsistente...',
            3: 'Controles de acceso a la información confidencial y crítica están bien implementados y se utilizan regularmente...',
            4: 'Los controles de acceso a la información confidencial y crítica cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de controles de acceso a la información confidencial y crítica es avanzada y supera los requisitos estándar...'
        },
        '¿Se establecen y mantienen políticas para la clasificación y etiquetado de la información dentro del departamento de sistemas?': {
            1: 'No se establecen ni mantienen políticas para la clasificación y etiquetado de la información...',
            2: 'Las políticas de clasificación y etiquetado existen pero no se mantienen adecuadamente...',
            3: 'Las políticas de clasificación y etiquetado están bien implementadas y se mantienen regularmente...',
            4: 'Las políticas de clasificación y etiquetado cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de políticas de clasificación y etiquetado es avanzada y supera los requisitos estándar...'
        }
    },
    'Gestión de Incidentes de Seguridad de la Información': {
        '¿Existe un procedimiento documentado para la gestión de incidentes de seguridad de la información?': {
            1: 'No hay un procedimiento documentado para la gestión de incidentes de seguridad...',
            2: 'El procedimiento para la gestión de incidentes existe pero no está actualizado o se implementa de manera limitada...',
            3: 'El procedimiento para la gestión de incidentes está bien documentado y se revisa regularmente...',
            4: 'El procedimiento para la gestión de incidentes cumple totalmente con los requisitos de ISO 27001...',
            5: 'La implementación del procedimiento para la gestión de incidentes es avanzada y supera los requisitos estándar...'
        },
        '¿Se realiza capacitación y simulacros periódicos para el personal sobre cómo responder a incidentes de seguridad de la información?': {
            1: 'No se realizan capacitaciones ni simulacros sobre incidentes de seguridad...',
            2: 'Las capacitaciones y simulacros se realizan de manera irregular o insuficiente...',
            3: 'Las capacitaciones y simulacros se realizan regularmente, cumpliendo con la mayoría de los requisitos de ISO 27001...',
            4: 'Las capacitaciones y simulacros cumplen totalmente con los requisitos de ISO 27001...',
            5: 'La implementación de capacitaciones y simulacros es avanzada y supera los requisitos estándar...'
        }
    }
}

# Procesar las calificaciones y calcular los promedios
def procesar_calificaciones(calificaciones):
    promedios = {aspecto: sum(valores[1] for valores in lista) / len(lista) for aspecto, lista in calificaciones.items()}
    promedios_ponderados = {aspecto: (promedio / 5) * 20 for aspecto, promedio in promedios.items()}

    calificacion_final = sum(promedios_ponderados.values()) / len(promedios_ponderados) * 5
    return promedios_ponderados, calificacion_final

# Generar gráfico de barras utilizando matplotlib
def generar_grafico(promedios_ponderados):
    aspectos = list(promedios_ponderados.keys())
    valores = list(promedios_ponderados.values())

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh(aspectos, valores, color='skyblue')
    ax.set_xlabel('Nivel de Cumplimiento (sobre 20)')
    ax.set_title('Gráfico de Nivel de Cumplimiento por Aspecto')
    ax.set_xlim(0, 20)

    st.pyplot(fig)

# Generar gráfico de radar utilizando matplotlib
def generar_grafico_radar(promedios_ponderados):
    etiquetas = list(promedios_ponderados.keys())
    valores = list(promedios_ponderados.values())
    valores += valores[:1]  # Añadir el primer valor al final para cerrar el gráfico

    angulos = np.linspace(0, 2 * np.pi, len(etiquetas), endpoint=False).tolist()
    angulos += angulos[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angulos, valores, color='skyblue', alpha=0.25)
    ax.plot(angulos, valores, color='skyblue', linewidth=2)
    ax.set_yticklabels([])
    ax.set_xticks(angulos[:-1])
    ax.set_xticklabels(etiquetas)
    ax.set_title('Gráfico de Radar por Aspecto')

    st.pyplot(fig)

# Generar la conclusión general basada en la calificación final
def generar_conclusion(calificacion_final):
    if 0 <= calificacion_final <= 25:
        return ("El departamento de sistemas muestra una falta significativa de cumplimiento en la gestión de acceso, "
                "seguridad física y ambiental, gestión de comunicaciones y operaciones, control de acceso a la información, "
                "y gestión de incidentes de seguridad de la información...")
    elif 26 <= calificacion_final <= 50:
        return ("El departamento de sistemas tiene algunos controles y políticas en su lugar, pero estos no son suficientemente robustos "
                "o no se aplican consistentemente...")
    elif 51 <= calificacion_final <= 75:
        return ("El departamento de sistemas ha implementado la mayoría de los controles de seguridad requeridos por la norma ISO 27001...")
    elif 76 <= calificacion_final <= 100:
        return ("El departamento de sistemas cumple completamente con los requisitos de la norma ISO 27001, y además implementa medidas adicionales...")
    else:
        return "Calificación no válida."

# Generar el informe en Word
def generar_informe_word(calificaciones, promedios_ponderados, calificacion_final, nombre_auditor, nombre_compania, fecha_evaluacion, destinatario, mensaje):
    document = Document()

    # Carátula
    document.add_heading('Informe de Evaluación de Cumplimiento de la Norma ISO 27001 (Sistema de Gestión de Seguridad de la Información)', 0)
    document.add_paragraph(f'Compañía Auditora: {nombre_compania}', style='Title')
    document.add_paragraph(f'Auditor: {nombre_auditor}', style='Heading 3')
    document.add_paragraph(f'Fecha de Evaluación: {fecha_evaluacion}', style='Heading 3')

    # Carta de introducción
    document.add_heading('Carta de Introducción', level=1)
    document.add_paragraph(f'Destinatario: {destinatario}', style='Heading 2')
    document.add_paragraph(mensaje)

    # Descripción del objetivo de la norma
    document.add_heading('Objetivo de la Norma ISO 27001', level=1)
    document.add_paragraph(
        "La norma ISO/IEC 27001 establece los requisitos para un sistema de gestión de seguridad de la información (SGSI)..."
    )

    # Descripción de las dimensiones evaluadas
    document.add_heading('Dimensiones Evaluadas', level=1)
    document.add_paragraph(
        "A continuación se detallan las diferentes dimensiones evaluadas en este informe, junto con una breve descripción de cada una:"
    )

    dimensiones = {
        'Gestión de Acceso': "Evalúa la existencia y eficacia de políticas y procedimientos para la gestión de accesos...",
        'Seguridad Física y Ambiental': "Evalúa las medidas de seguridad física y controles ambientales implementados para proteger...",
        'Gestión de Comunicaciones y Operaciones': "Evalúa los procedimientos seguros para la transmisión de datos sensibles y las prácticas de gestión de operaciones...",
        'Control de Acceso a la Información': "Evalúa los controles implementados para limitar el acceso a la información confidencial y crítica...",
        'Gestión de Incidentes de Seguridad de la Información': "Evalúa la existencia y eficacia de procedimientos para la gestión de incidentes..."
    }

    for dimension, descripcion in dimensiones.items():
        document.add_heading(dimension, level=2)
        document.add_paragraph(descripcion)

    # Metodología de calificación
    document.add_heading('Metodología de Calificación', level=1)
    document.add_paragraph(
        "La evaluación se basa en una escala de 1 a 5, donde cada valor representa el nivel de cumplimiento de la norma:"
    )
    calificacion_metodologia = {
        1: "1 = No Cumple: No se realiza ninguna acción o la acción es insuficiente.",
        2: "2 = Cumple Parcialmente: Las acciones se realizan pero no con la frecuencia o efectividad requerida.",
        3: "3 = Cumple en Gran Medida: Las acciones se realizan regularmente y cumplen con la mayoría de los requisitos.",
        4: "4 = Cumple Totalmente: Las acciones cumplen con todos los requisitos establecidos.",
        5: "5 = Cumple y Supera las Expectativas: Se implementan medidas adicionales que superan los requisitos establecidos."
    }

    for key, value in calificacion_metodologia.items():
        document.add_paragraph(value)

    # Resultados de la evaluación
    document.add_heading('Resultados de la Evaluación', level=1)
    for aspecto, preguntas in calificaciones.items():
        document.add_heading(aspecto, level=2)
        for pregunta, calificacion in preguntas:
            descripcion = rubricas[aspecto][pregunta][calificacion]
            p = document.add_paragraph()
            p.add_run(f'{pregunta}: ').bold = True
            p.add_run(f'{calificacion} - {descripcion}')
        document.add_paragraph(f'Promedio del aspecto ({aspecto}): {promedios_ponderados[aspecto]:.2f} / 20')
        document.add_paragraph()

    document.add_paragraph(f'Calificación final del departamento de sistemas: {calificacion_final:.2f} / 100')
    document.add_paragraph()

    # Conclusión general
    conclusion = generar_conclusion(calificacion_final)
    document.add_heading('Conclusión General', level=1)
    document.add_paragraph(conclusion)
    document.add_paragraph()

    # Añadir gráfico de barras
    document.add_heading('Gráfico de Nivel de Cumplimiento por Aspecto', level=1)
    generar_grafico(promedios_ponderados)
    document.add_picture('grafico_cumplimiento.png', width=Inches(6))

    # Añadir gráfico de radar
    document.add_heading('Gráfico de Radar por Aspecto', level=1)
    generar_grafico_radar(promedios_ponderados)
    document.add_picture('grafico_radar.png', width=Inches(6))

    # Añadir pie de página
    section = document.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f'Compañía Auditora: {nombre_compania} - Fecha de Evaluación: {fecha_evaluacion}'
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save('informe.docx')
    st.success("El informe se ha generado correctamente en informe.docx")

# Interfaz con Streamlit
def main():
    st.title("Evaluación de Cumplimiento ISO 27001")
    calificaciones = {key: [] for key in rubricas.keys()}

    for aspecto in rubricas.keys():
        st.header(aspecto)
        for pregunta in rubricas[aspecto].keys():
            calificacion = st.selectbox(pregunta, options=[f'{i}: {rubricas[aspecto][pregunta][i]}' for i in range(1, 6)], key=f"{aspecto}_{pregunta}")
            calificaciones[aspecto].append((pregunta, int(calificacion.split(":")[0])))

    if st.button("Generar Informe"):
        nombre_auditor = st.text_input("Nombre del Auditor")
        nombre_compania = st.text_input("Nombre de la Compañía")
        fecha_evaluacion = st.text_input("Fecha de Evaluación (DD/MM/AAAA)", value=str(datetime.today().strftime('%d/%m/%Y')))
        destinatario = st.text_input("Destinatario del Informe")
        mensaje = st.text_area("Carta de Introducción")

        if not all([nombre_auditor, nombre_compania, fecha_evaluacion, destinatario, mensaje]):
            st.error("Debe completar todos los campos para generar el informe.")
        else:
            promedios_ponderados, calificacion_final = procesar_calificaciones(calificaciones)
            generar_informe_word(calificaciones, promedios_ponderados, calificacion_final, nombre_auditor, nombre_compania, fecha_evaluacion, destinatario, mensaje)

            st.write("**Calificación Final:**", calificacion_final)
            st.write("**Conclusión:**", generar_conclusion(calificacion_final))
            generar_grafico(promedios_ponderados)
            generar_grafico_radar(promedios_ponderados)

if __name__ == "__main__":
    main()
